"""
One-off: convert docs/PRESENTATION-EN.md -> docs/PRESENTATION-EN.docx (Word).
Run: python docs/build_presentation_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "PRESENTATION-EN.md"
OUT_PATH = ROOT / "PRESENTATION-EN.docx"


def is_table_separator(line: str) -> bool:
    cells = [c.strip() for c in line.strip().split("|")[1:-1]]
    if not cells:
        return False
    return all(re.fullmatch(r":?-{3,}:?", c) for c in cells)


def parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().split("|")[1:-1]]


def strip_md_inline_for_table(cell: str) -> str:
    cell = re.sub(r"\*\*(.+?)\*\*", r"\1", cell)
    cell = re.sub(r"`([^`]+)`", r"\1", cell)
    return cell


def add_runs_from_inline(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    i = 0
    n = len(text)
    while i < n:
        if text.startswith("**", i):
            j = text.find("**", i + 2)
            if j != -1:
                run = p.add_run(text[i + 2 : j])
                run.bold = True
                i = j + 2
                continue
        if text.startswith("`", i):
            j = text.find("`", i + 1)
            if j != -1:
                run = p.add_run(text[i + 1 : j])
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                i = j + 1
                continue
        j_star = text.find("**", i)
        j_bt = text.find("`", i)
        candidates = [x for x in (j_star, j_bt) if x != -1]
        j = min(candidates) if candidates else n
        p.add_run(text[i:j])
        i = j


def main() -> None:
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    doc = Document()
    title = lines[0].lstrip("# ").strip() if lines else "Presentation"
    h = doc.add_heading(title, 0)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    i = 1
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue
        if stripped == "---":
            i += 1
            continue

        # Markdown table block
        if stripped.startswith("|") and not is_table_separator(stripped):
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                if is_table_separator(row_line):
                    i += 1
                    continue
                rows.append([strip_md_inline_for_table(c) for c in parse_table_row(row_line)])
                i += 1
            if rows:
                ncol = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=ncol)
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci in range(ncol):
                        cell_text = row[ci] if ci < len(row) else ""
                        table.rows[ri].cells[ci].text = cell_text
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            add_runs_from_inline(doc, text, style="List Number")
            i += 1
            continue

        if stripped.startswith("- "):
            add_runs_from_inline(doc, stripped[2:], style="List Bullet")
            i += 1
            continue

        add_runs_from_inline(doc, stripped)
        i += 1

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
